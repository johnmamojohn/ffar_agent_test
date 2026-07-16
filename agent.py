"""
FFAR Partnership Research Agent — production module.

This module contains the core agent logic, separated from the UI layer.
The Streamlit app (app.py) imports functions from here.
"""

import os
from typing import List
from pydantic import BaseModel, Field

from langchain_anthropic import ChatAnthropic
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_community.document_loaders import Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.tools import create_retriever_tool
from langchain_core.messages import HumanMessage
from langgraph.prebuilt import create_react_agent

from docx import Document


# ============================================================
# SCHEMA
# ============================================================

class KeyFacts(BaseModel):
    organization_name: str = Field(description="Official organization name")
    headquarters: str = Field(description="City, Country of headquarters")
    structure: str = Field(description="Public company, Private, Nonprofit, Foundation, Government agency, etc.")
    fiscal_year_end: str = Field(description="Fiscal year end (e.g., 'December 31') or 'Not publicly available'")
    geographic_focus: str = Field(description="Geographic areas of operation")
    core_mission: str = Field(description="One-sentence mission statement")
    business_offerings: str = Field(description="Main products or services")


class GeneralBusinessProfile(BaseModel):
    key_facts: KeyFacts
    publicly_stated_goals: List[str] = Field(description="Sustainability, research, or strategic commitments with time periods")
    existing_partnerships: List[str] = Field(description="Known R&D partnerships in food & agriculture, with years if available")
    ffar_opportunity_analysis: str = Field(description="General assessment of alignment with FFAR's mission areas")


# ============================================================
# AGENT BUILDER
# ============================================================

def build_agent_pipeline(template_path: str = "data/Business Profile Template.docx"):
    """
    Build the full agent pipeline: LLM + tools + RAG + structured output.
    Returns a research function ready to call.
    """
    # LLM
    llm = ChatAnthropic(
        model="claude-haiku-4-5-20251001",
        temperature=0
    )
    
    # Web search tool
    web_search = TavilySearchResults(max_results=5)
    
    # RAG: load, split, embed, store the FFAR template
    loader = Docx2txtLoader(template_path)
    documents = loader.load()
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    chunks = splitter.split_documents(documents)
    
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name="ffar_template"
    )
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    
    # Wrap retriever as a tool
    ffar_template_tool = create_retriever_tool(
        retriever=retriever,
        name="ffar_template_lookup",
        description=(
            "Use this tool to look up the FFAR Business Profile Template structure. "
            "Call this before writing the final answer to ensure the output follows "
            "the FFAR template format."
        )
    )
    
    tools = [web_search, ffar_template_tool]
    
    system_prompt = """You are a Partnership Research Analyst for FFAR (Foundation for Food & Agriculture Research).

You have two tools:
1. web_search — for recent news and public information
2. ffar_template_lookup — for the FFAR Business Profile Template structure

Workflow:
1. Use ffar_template_lookup to understand the template structure.
2. Use web_search to gather organization information.
3. Compose your answer following the FFAR template format.

Be factual. If information is not publicly available, say so — never fabricate.
"""
    
    rag_agent = create_react_agent(
        model=llm,
        tools=tools,
        prompt=system_prompt
    )
    
    structured_llm = llm.with_structured_output(GeneralBusinessProfile)
    
    return rag_agent, structured_llm


# ============================================================
# RESEARCH FUNCTION
# ============================================================

def research_organization(organization_name: str, rag_agent, structured_llm) -> GeneralBusinessProfile:
    """
    Full research pipeline for a single organization.
    """
    research_task = f"Research {organization_name} thoroughly for a FFAR business profile. Use both the FFAR template lookup and web search."
    
    research_response = rag_agent.invoke({
        "messages": [HumanMessage(content=research_task)]
    })
    research_text = research_response["messages"][-1].content
    
    extraction_prompt = f"""Based on the following research about {organization_name}, extract the information into the required structured format.

Research:
{research_text}

Extract all relevant fields. For missing information, use "Not publicly available"."""
    
    return structured_llm.invoke(extraction_prompt)


# ============================================================
# WORD EXPORT
# ============================================================

def profile_to_word(profile: GeneralBusinessProfile, output_path: str) -> str:
    """Export a profile to a Word document following FFAR template."""
    doc = Document()
    
    doc.add_heading("Business Profile & Partnership Opportunity", level=0)
    doc.add_heading(profile.key_facts.organization_name, level=1)
    
    # Key Facts table
    doc.add_heading("Key Facts", level=2)
    facts_table = doc.add_table(rows=7, cols=2)
    facts_table.style = "Light Grid Accent 1"
    
    facts = [
        ("Organization Name", profile.key_facts.organization_name),
        ("Headquarters", profile.key_facts.headquarters),
        ("Structure", profile.key_facts.structure),
        ("Fiscal Year End", profile.key_facts.fiscal_year_end),
        ("Geographic Focus", profile.key_facts.geographic_focus),
        ("Core Mission", profile.key_facts.core_mission),
        ("Business Offerings", profile.key_facts.business_offerings),
    ]
    for i, (label, value) in enumerate(facts):
        facts_table.cell(i, 0).text = label
        facts_table.cell(i, 1).text = value
    
    doc.add_heading("Publicly Stated Goals & Commitments", level=2)
    for goal in profile.publicly_stated_goals:
        doc.add_paragraph(goal, style="List Bullet")
    
    doc.add_heading("Existing R&D Partnerships in Food & Agriculture", level=2)
    for partnership in profile.existing_partnerships:
        doc.add_paragraph(partnership, style="List Bullet")
    
    doc.add_heading("FFAR Opportunity Analysis", level=2)
    doc.add_paragraph(profile.ffar_opportunity_analysis)
    
    doc.save(output_path)
    return output_path